"""
Geração de contratos de administração de obra em formato .docx, a partir
dos dados cadastrados em Contratos_ADM (via GestaoContratos). Suporta os
três métodos de pagamento: Percentual da Quinzena, Valor Fixo em Parcelas,
e Eventos/Fases.

Extraído de Sistema_Entrada_Dados.py em [DATA_DA_EXTRACAO].

CORREÇÃO aplicada nesta extração (não é mudança de lógica de negócio,
é correção de um import quebrado):
    Em _obter_indice_correcao_padrao, o import original era:
        from correcao_monetaria import GerenciadorCorrecaoMonetaria
    Isso está incorreto — o módulo real é src/correcao_monetaria.py, e o
    import sem o prefixo "src." falhava sempre (import absoluto sem o
    pacote correto), fazendo o método cair silenciosamente no fallback
    'IGPM', independente do índice de correção configurado pelo cliente
    (ex: INCC, IPCA). Corrigido abaixo para:
        from src.correcao_monetaria import GerenciadorCorrecaoMonetaria
    Isso significa que, a partir desta extração, contratos gerados com
    "Valor Fixo em Parcelas" ou "Eventos/Fases" (que citam o índice de
    correção na Cláusula Sétima) podem passar a exibir um índice
    diferente de 'IGPM' se o sistema estiver configurado assim. Vale
    conferir o primeiro contrato gerado após a correção para confirmar
    que o índice exibido está correto.
"""
import re
from datetime import datetime
from pathlib import Path

from openpyxl import load_workbook


class GeradorContratoADM:
    """
    Gera contratos de administração de obra em formato DOCX.

    Suporta os três métodos de pagamento cadastrados em Contratos_ADM:
      • "Percentual da Quinzena"   → remunera % sobre o movimento quinzenal
      • "Valor Fixo em Parcelas"   → honorário fixo total em N parcelas com
                                     datas de vencimento definidas
      • "Eventos/Fases"            → honorário vinculado a marcos da obra

    Uso
    ---
    gerador = GeradorContratoADM()
    paths = gerador.gerar_contratos_do_contrato(
                nome_cliente="CLIENTE EXEMPLO",
                num_contrato="2024/08",
                arquivo_cliente=PASTA_CLIENTES / "CLIENTE_EXEMPLO.xlsx",
                arquivo_clientes_geral=ARQUIVO_CLIENTES,
                arquivo_fornecedores=ARQUIVO_FORNECEDORES,
    )
    # paths → lista de caminhos gerados (um por administrador)
    """

    PASTA_CONTRATOS_ADM = None   # definir como PASTA_CLIENTES / "Contratos_ADM"

    # ── utilidades ──────────────────────────────────────────────

    @staticmethod
    def _formatar_doc(numero: str) -> str:
        """Formata CPF (11 dígitos) ou CNPJ (14 dígitos)."""
        d = re.sub(r'\D', '', str(numero))
        if len(d) == 11:
            return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
        if len(d) == 14:
            return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}"
        return str(numero)

    @staticmethod
    def _tipo_pessoa(cnpj_cpf: str) -> str:
        """Retorna 'PF' ou 'PJ' baseado no número de dígitos."""
        d = re.sub(r'\D', '', str(cnpj_cpf))
        return 'PF' if len(d) == 11 else 'PJ'

    @staticmethod
    def _extenso(valor: float) -> str:
        """Valor monetário por extenso (requer num2words)."""
        try:
            from num2words import num2words
            return num2words(valor, lang='pt_BR', to='currency')
        except Exception as e:
            # Antes só capturava ImportError e falhava silenciosamente —
            # agora loga o motivo real (ImportError, dado ausente no
            # bundle do PyInstaller, versão incompatível, etc.) para não
            # gerar contratos "errados" sem nenhum rastro no log.
            import logging
            logging.getLogger("sistema").debug(
                f"num2words indisponível para valor por extenso: "
                f"{type(e).__name__}: {e}"
            )
            return f"{valor:.2f} reais"

    @staticmethod
    def _data_extenso(data_obj) -> str:
        meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
                 'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
        if isinstance(data_obj, str):
            data_obj = datetime.strptime(data_obj, '%d/%m/%Y')
        return f"{data_obj.day} de {meses[data_obj.month - 1]} de {data_obj.year}"

    @staticmethod
    def _fmt_valor(valor: float) -> str:
        """Formata valor monetário em reais BR."""
        return f"R$ {valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

    @staticmethod
    def _fmt_data(data_obj) -> str:
        """Formata data como dd/mm/aaaa."""
        if data_obj is None:
            return ''
        if isinstance(data_obj, str):
            return data_obj
        return data_obj.strftime('%d/%m/%Y')

    # ── leitura de dados ────────────────────────────────────────

    def _ler_dados_contrato(self, arquivo_cliente: Path, num_contrato: str) -> dict:
        """
        Lê Contratos_ADM e retorna um dict com todos os dados do
        contrato e de seus administradores.
        """
        wb = load_workbook(arquivo_cliente, data_only=True)
        ws = wb['Contratos_ADM']

        contrato = {
            'num_contrato': num_contrato,
            'data_inicio': None,
            'data_fim': None,
            'status': None,
            'observacoes': None,
            'valor_global': 0.0,
            'admins': [],
            'parcelas': [],
            # índice de correção cadastrado no contrato (coluna futura ou observação)
            'indice_correcao': None,
        }

        # leitura do cabeçalho do contrato (colunas A-F)
        for row in ws.iter_rows(min_row=3, values_only=True):
            if str(row[0] or '').strip() == str(num_contrato).strip():
                contrato['data_inicio']   = row[1]
                contrato['data_fim']      = row[2]
                contrato['status']        = row[3]
                contrato['observacoes']   = row[4]
                contrato['valor_global']  = float(row[5] or 0)
                # coluna F+1 (índice 5) = valor_global, coluna G+ seria extra
                # Se houver coluna para índice de correção, ler aqui:
                # contrato['indice_correcao'] = row[6] (ajustar conforme estrutura)
                break

        # leitura dos administradores (coluna G = num_contrato do admin)
        admins_vistos = set()
        for row in ws.iter_rows(min_row=3, values_only=True):
            if str(row[6] or '').strip() == str(num_contrato).strip():
                cnpj = str(row[7] or '').strip()
                if cnpj and cnpj not in admins_vistos:
                    admins_vistos.add(cnpj)
                    try:
                        valor_total = float(str(row[11] or 0).replace(',', '.'))
                    except ValueError:
                        valor_total = 0.0
                    try:
                        perc_raw = str(row[10] or '').replace('%', '').replace(',', '.')
                        percentual = float(perc_raw) if perc_raw else 0.0
                    except ValueError:
                        percentual = 0.0
                    contrato['admins'].append({
                        'cnpj_cpf':    cnpj,
                        'nome':        str(row[8] or '').strip(),
                        'tipo':        str(row[9] or '').strip(),
                        'percentual':  percentual,
                        'valor_total': valor_total,
                        'num_parcelas': int(row[12] or 0),
                        'tipo_pessoa': self._tipo_pessoa(cnpj),
                    })

        # leitura das parcelas/eventos (coluna Y = referência)
        for row in ws.iter_rows(min_row=3, values_only=True):
            if str(row[24] or '').strip() == str(num_contrato).strip():
                try:
                    valor_parc = float(str(row[29] or 0).replace(',', '.'))
                except (ValueError, TypeError):
                    valor_parc = 0.0

                perc_raw = row[33]
                if perc_raw is not None:
                    try:
                        percentual_ev = float(
                            str(perc_raw).replace('%', '').replace(',', '.').strip()
                        )
                    except (ValueError, TypeError):
                        percentual_ev = 0.0
                else:
                    percentual_ev = 0.0

                # Número da parcela: 0 = entrada/sinal
                num_parcela = row[25]

                contrato['parcelas'].append({
                    'numero':     num_parcela,
                    'cnpj_cpf':   str(row[26] or '').strip(),
                    'nome':       str(row[27] or '').strip(),
                    'vencimento': row[28],   # data de vencimento
                    'valor':      valor_parc,
                    'status':     str(row[30] or '').strip(),
                    'descricao':  str(row[32] or '').strip(),
                    'percentual': percentual_ev,
                    'eh_entrada': (num_parcela == 0),
                })

        # Recalcular percentuais ausentes
        for parcela in contrato['parcelas']:
            if parcela['percentual'] == 0.0 and parcela['valor'] > 0:
                cnpj_p = parcela['cnpj_cpf']
                admin_match = next(
                    (a for a in contrato['admins']
                     if re.sub(r'\D', '', str(a['cnpj_cpf'])) ==
                        re.sub(r'\D', '', str(cnpj_p))),
                    None
                )
                if admin_match and admin_match['valor_total'] > 0:
                    parcela['percentual'] = parcela['valor'] / admin_match['valor_total']

        wb.close()
        return contrato

    def _ler_dados_cliente(self, arquivo_clientes_geral: Path, nome_cliente: str) -> dict:
        """Busca dados do contratante em Clientes.xlsx."""
        import pandas as pd
        df = pd.read_excel(arquivo_clientes_geral)
        row = df[df['Nome'] == nome_cliente]
        if row.empty:
            return {'nome': nome_cliente, 'cpf': '', 'cno': '',
                    'estado_civil': '', 'endereco': '', 'cidade': 'Belo Horizonte',
                    'estado': 'MG'}
        r = row.iloc[0]

        def _safe(col, default=''):
            import math
            v = r.get(col, default)
            return default if (v is None or (isinstance(v, float) and math.isnan(v))) else str(v)

        cpf_raw = r.get('CPF', '')
        try:
            cpf_raw = str(int(float(cpf_raw)))
        except (ValueError, TypeError):
            cpf_raw = str(cpf_raw or '')

        return {
            'nome':         _safe('Nome'),
            'cpf':          self._formatar_doc(cpf_raw),
            'cno':          _safe('CNO'),
            'estado_civil': _safe('Estado Civil', 'não informado'),
            'endereco':     _safe('Endereço', 'não informado'),
            'cidade':       _safe('Cidade', 'Belo Horizonte'),
            'estado':       _safe('Estado', 'MG'),
        }

    def _ler_dados_fornecedor(self, arquivo_fornecedores: Path, cnpj_cpf: str) -> dict:
        """Busca dados do administrador (contratado) em base_fornecedores.xlsx."""
        import pandas as pd
        import math

        def _safe_val(r, *cols, default=''):
            for col in cols:
                v = r.get(col, None)
                if v is not None and not (isinstance(v, float) and math.isnan(v)):
                    return str(v).strip()
            return default

        try:
            df = pd.read_excel(arquivo_fornecedores)
        except Exception:
            return {'nome': 'não informado', 'cnpj_cpf': self._formatar_doc(cnpj_cpf),
                    'endereco': 'não informado', 'dados_bancarios': ''}

        col_doc = None
        for candidato in ('CNPJ/CPF', 'CPF', 'CNPJ', 'cnpj_cpf', 'documento'):
            if candidato in df.columns:
                col_doc = candidato
                break

        if col_doc is None:
            return {'nome': 'não informado', 'cnpj_cpf': self._formatar_doc(cnpj_cpf),
                    'endereco': 'não informado', 'dados_bancarios': ''}

        cnpj_limpo = re.sub(r'\D', '', cnpj_cpf)
        df['_cnpj_limpo'] = df[col_doc].apply(
            lambda x: re.sub(r'\D', '', str(x) if x is not None and str(x) != 'nan' else ''))
        row = df[df['_cnpj_limpo'] == cnpj_limpo]

        if row.empty:
            return {'nome': 'não informado', 'cnpj_cpf': self._formatar_doc(cnpj_cpf),
                    'endereco': 'não informado', 'dados_bancarios': ''}

        r = row.iloc[0]
        dados_banc = _safe_val(r, 'DADOS BANCÁRIOS', 'dados_bancarios', 'PIX', 'pix')
        return {
            'nome':            _safe_val(r, 'NOME', 'Nome', 'nome', default='não informado'),
            'cnpj_cpf':        self._formatar_doc(cnpj_cpf),
            'endereco':        _safe_val(r, 'ENDEREÇO', 'Endereço', 'ENDERECO',
                                         'endereco', default='não informado'),
            'dados_bancarios': dados_banc,
            'razao_social':    _safe_val(r, 'RAZÃO SOCIAL', 'NOME', 'Nome',
                                         default='não informado'),
        }

    # ── obter índice de correção configurado ────────────────────

    @staticmethod
    def _obter_indice_correcao_padrao() -> str:
        """
        Retorna o índice de correção monetária padrão configurado no sistema
        (GerenciadorCorrecaoMonetaria → parametros_sistema.json).
        Retorna 'IGPM' como fallback.
        """
        try:
            # CORRIGIDO nesta extração: faltava o prefixo "src." — ver
            # docstring do módulo para detalhes do impacto.
            from src.correcao_monetaria import GerenciadorCorrecaoMonetaria
            gcm = GerenciadorCorrecaoMonetaria()
            return gcm.config.get('indices_correcao', {}).get('indice_padrao', 'IGPM')
        except Exception:
            return 'IGPM'

    # ── construção do documento Word ────────────────────────────

    def _build_doc(self, contrato: dict, admin: dict,
                   dados_cliente: dict, dados_fornecedor: dict,
                   metodo: str):
        """
        Constrói o Document python-docx para um par contratante/contratado.

        metodo : "Percentual da Quinzena" | "Valor Fixo em Parcelas" | "Eventos/Fases"
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── estilos ─────────────────────────────────────────────
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(11)
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_normal.paragraph_format.space_after = Pt(6)

        h1 = doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(12)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)

        # ── margens ─────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # ── helpers locais ───────────────────────────────────────
        def par(text='', bold=False, center=False, size=None,
                space_before=None, space_after=None):
            p = doc.add_paragraph()
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if space_before is not None:
                p.paragraph_format.space_before = Pt(space_before)
            if space_after is not None:
                p.paragraph_format.space_after = Pt(space_after)
            if text:
                r = p.add_run(text)
                r.bold = bold
                if size:
                    r.font.size = Pt(size)
            return p

        def heading(text):
            doc.add_heading(text, level=1)

        def _cell_text(cell, text, bold=False, size=9):
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)

        # ── dados formatados ─────────────────────────────────────
        nome_cliente    = dados_cliente['nome']
        cpf_cliente     = dados_cliente['cpf']
        cno_cliente     = dados_cliente['cno']
        ec_cliente      = dados_cliente['estado_civil']
        end_cliente     = dados_cliente['endereco']
        cidade          = dados_cliente['cidade']

        nome_adm        = dados_fornecedor['nome']
        doc_adm         = dados_fornecedor['cnpj_cpf']
        end_adm         = dados_fornecedor['endereco']
        dados_banc      = dados_fornecedor['dados_bancarios']
        tipo_pessoa_adm = admin['tipo_pessoa']

        data_contrato_obj = contrato['data_inicio']
        if isinstance(data_contrato_obj, str):
            data_contrato_obj = datetime.strptime(data_contrato_obj, '%d/%m/%Y')
        data_extenso   = self._data_extenso(data_contrato_obj)
        data_inicio_fmt = self._fmt_data(data_contrato_obj)
        data_fim_fmt    = self._fmt_data(contrato['data_fim'])

        valor_global   = admin['valor_total']
        multa_float    = valor_global * 0.10
        valor_fmt      = self._fmt_valor(valor_global)
        multa_fmt      = self._fmt_valor(multa_float)
        valor_extenso  = self._extenso(valor_global)
        multa_extenso  = self._extenso(multa_float)

        # índice de correção
        indice_correcao = (contrato.get('indice_correcao')
                           or self._obter_indice_correcao_padrao())

        if tipo_pessoa_adm == 'PJ':
            qualif_adm     = (f"pessoa jurídica inscrita no CNPJ sob o n.º {doc_adm}, "
                              f"com sede na {end_adm}")
            denominacao_adm = "CONTRATADA"
        else:
            qualif_adm     = (f"pessoa física inscrita no CPF sob o n.º {doc_adm}, "
                              f"residente na {end_adm}")
            denominacao_adm = "CONTRATADO"

        # ── TÍTULO ───────────────────────────────────────────────
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tit = titulo.add_run(
            "CONTRATO PARTICULAR DE PRESTAÇÃO DE SERVIÇOS DE "
            "ADMINISTRAÇÃO DE OBRA")
        run_tit.bold = True
        run_tit.font.size = Pt(13)

        subtit = doc.add_paragraph()
        subtit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtit.add_run(
            f"Contrato n.º {contrato['num_contrato']}"
        ).font.size = Pt(10)

        # ── PREÂMBULO ────────────────────────────────────────────
        p_intro = doc.add_paragraph()
        p_intro.add_run(
            f"Aos {data_extenso}, nesta cidade de {cidade}, "
            f"entre as partes abaixo identificadas:")

        p_cliente = doc.add_paragraph()
        p_cliente.add_run(nome_cliente).bold = True
        p_cliente.add_run(
            f", pessoa física inscrita no CNO n.º {cno_cliente} "
            f"e CPF n.º {cpf_cliente}, {ec_cliente}, "
            f"residente e domiciliado na {end_cliente}, "
            f"doravante denominado(a) simplesmente CONTRATANTE; e")

        p_adm = doc.add_paragraph()
        p_adm.add_run(nome_adm).bold = True
        p_adm.add_run(
            f", {qualif_adm}, "
            f"doravante denominado(a) {denominacao_adm}; "
            f"têm entre si, justo e contratado o seguinte:")

        # ── CLÁUSULA PRIMEIRA — OBJETO ───────────────────────────
        heading("CLÁUSULA PRIMEIRA — OBJETO")
        doc.add_paragraph(
            "O presente contrato tem por objeto a prestação de serviços "
            "de administração e gerenciamento da execução da obra de "
            f"construção civil localizada no endereço: {end_cliente}, "
            "compreendendo as atividades de coordenação técnica, "
            "fiscalização, controle financeiro, supervisão de equipes e "
            "demais atividades inerentes à função de gestor de obras.")

        p1 = doc.add_paragraph()
        p1.add_run("PARÁGRAFO ÚNICO: ").bold = True
        p1.add_run(
            "São atribuições do(a) CONTRATADO(A), sem caráter exaustivo: "
            "coordenar e fiscalizar a execução dos serviços; supervisionar "
            "equipes e subcontratados; gerenciar o cronograma físico-"
            "financeiro; elaborar relatórios quinzenais de prestação de "
            "contas; realizar cotações e aprovar compras de materiais; "
            "zelar pelo cumprimento das normas técnicas e de segurança.")

        # ── CLÁUSULA SEGUNDA — PRAZO ─────────────────────────────
        heading("CLÁUSULA SEGUNDA — PRAZO")
        doc.add_paragraph(
            f"Os serviços terão início em {data_inicio_fmt} e "
            f"previsão de término em {data_fim_fmt}, podendo ser "
            "prorrogados mediante Termo Aditivo assinado pelas partes, "
            "sem que isso implique em qualquer majoração automática da "
            "remuneração, salvo acordo expresso.")

        # ── CLÁUSULA TERCEIRA — REMUNERAÇÃO (varia por método) ───
        heading("CLÁUSULA TERCEIRA — REMUNERAÇÃO")

        # ─── Parcelas do administrador (filtradas)
        cnpj_admin_limpo = re.sub(r'\D', '', str(admin['cnpj_cpf']))
        parcelas_admin = [
            p for p in contrato['parcelas']
            if re.sub(r'\D', '', str(p['cnpj_cpf'])) == cnpj_admin_limpo
        ]
        # Separar entrada das demais
        entrada = next((p for p in parcelas_admin if p.get('eh_entrada')), None)
        parcelas_comuns = [p for p in parcelas_admin if not p.get('eh_entrada')]

        if metodo == "Percentual da Quinzena":
            # ── Percentual da Quinzena ──────────────────────────
            perc = admin['percentual']

            doc.add_paragraph(
                f"Como remuneração pelos serviços prestados, o(a) "
                f"CONTRATANTE pagará ao(à) {denominacao_adm} o percentual "
                f"de {perc:.2f}% ({self._extenso_percentual(perc)} por cento) "
                "incidente sobre o total de pagamentos realizados em cada "
                "quinzena de referência do relatório de obra.")

            p_q1 = doc.add_paragraph()
            p_q1.add_run("PARÁGRAFO PRIMEIRO: ").bold = True
            p_q1.add_run(
                "A base de cálculo corresponde à soma dos lançamentos "
                "dos tipos 1 a 6 registrados no relatório de cada "
                "quinzena (dias 5 e 20 de cada mês), excluídos os "
                "próprios valores de administração.")

            p_q2 = doc.add_paragraph()
            p_q2.add_run("PARÁGRAFO SEGUNDO: ").bold = True
            p_q2.add_run(
                "O pagamento será efetuado até o vencimento constante "
                "no lançamento quinzenal correspondente, mediante "
                "transferência bancária ou PIX nos dados informados "
                "ao final deste instrumento.")

            if tipo_pessoa_adm == 'PJ':
                p_nf = doc.add_paragraph()
                p_nf.add_run("PARÁGRAFO TERCEIRO — NOTA FISCAL: ").bold = True
                p_nf.add_run(
                    "O(A) CONTRATADO(A), por ser pessoa jurídica, "
                    "deverá emitir Nota Fiscal de Serviços (NFS-e) "
                    "correspondente a cada pagamento quinzenal, sob "
                    "pena de retenção do valor até a regularização fiscal.")

        elif metodo == "Valor Fixo em Parcelas":
            # ── Valor Fixo em Parcelas ──────────────────────────
            num_parc = admin['num_parcelas'] or len(parcelas_comuns)

            # Texto introdutório
            if entrada:
                valor_entrada_fmt = self._fmt_valor(entrada['valor'])
                valor_restante    = valor_global - entrada['valor']
                valor_parc_fmt    = self._fmt_valor(valor_restante / num_parc
                                                    if num_parc else valor_restante)
                intro_text = (
                    f"Como remuneração pelos serviços prestados, o(a) "
                    f"CONTRATANTE pagará ao(à) {denominacao_adm} o honorário "
                    f"total de {valor_fmt} ({valor_extenso}), sendo "
                    f"{valor_entrada_fmt} a título de sinal/entrada, "
                    f"e o saldo restante dividido em {num_parc} parcela(s) "
                    f"de {valor_parc_fmt} cada, conforme cronograma abaixo:")
            else:
                valor_parcela = valor_global / num_parc if num_parc else valor_global
                valor_parc_fmt = self._fmt_valor(valor_parcela)
                intro_text = (
                    f"Como remuneração pelos serviços prestados, o(a) "
                    f"CONTRATANTE pagará ao(à) {denominacao_adm} o honorário "
                    f"total de {valor_fmt} ({valor_extenso}), dividido em "
                    f"{num_parc} parcela(s) de {valor_parc_fmt} cada, "
                    f"conforme cronograma abaixo:")

            doc.add_paragraph(intro_text)

            todas_parcelas_tabela = []
            if entrada:
                todas_parcelas_tabela.append(entrada)
            todas_parcelas_tabela.extend(
                sorted(parcelas_comuns, key=lambda x: x['numero'] or 0)
            )

            if todas_parcelas_tabela:
                tab = doc.add_table(rows=1, cols=4)
                tab.style = 'Table Grid'

                # larguras: Parcela | Descrição | Valor | Data Vencimento
                col_w = [Inches(0.8), Inches(3.5), Inches(1.1), Inches(1.1)]
                for i, w in enumerate(col_w):
                    for cell in tab.columns[i].cells:
                        cell.width = w

                # Largura total no XML
                tbl = tab._tbl
                tblPr = tbl.tblPr
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '9360')
                tblW.set(qn('w:type'), 'dxa')
                tblPr.append(tblW)

                hdr = tab.rows[0].cells
                _cell_text(hdr[0], 'Parcela',        bold=True)
                _cell_text(hdr[1], 'Descrição',       bold=True)
                _cell_text(hdr[2], 'Valor (R$)',       bold=True)
                _cell_text(hdr[3], 'Data Vencimento',  bold=True)

                total_tabela = 0.0
                for parc in todas_parcelas_tabela:
                    row_cells = tab.add_row().cells
                    for i, w in enumerate(col_w):
                        row_cells[i].width = w

                    if parc.get('eh_entrada') or parc['numero'] == 0:
                        num_display = 'ENTRADA'
                    else:
                        num_display = str(parc['numero'])

                    _total_lanc = len(todas_parcelas_tabela)
                    _pos = todas_parcelas_tabela.index(parc) + 1
                    desc = f"ADM. OBRA - PARCELA {_pos}/{_total_lanc}"
                    val_fmt = self._fmt_valor(parc['valor'])
                    dt_fmt  = self._fmt_data(parc.get('vencimento'))

                    _cell_text(row_cells[0], num_display)
                    _cell_text(row_cells[1], desc[:100] + ('…' if len(desc) > 100 else ''))
                    _cell_text(row_cells[2], val_fmt)
                    _cell_text(row_cells[3], dt_fmt)

                    total_tabela += parc['valor']

                row_total = tab.add_row().cells
                for i, w in enumerate(col_w):
                    row_total[i].width = w
                _cell_text(row_total[0], 'TOTAL', bold=True)
                _cell_text(row_total[1], '',      bold=True)
                _cell_text(row_total[2], self._fmt_valor(total_tabela), bold=True)
                _cell_text(row_total[3], '',      bold=True)

            par(space_before=8)

            p_f1 = doc.add_paragraph()
            p_f1.add_run("PARÁGRAFO PRIMEIRO: ").bold = True
            p_f1.add_run(
                "Os valores e as datas de vencimento constantes no "
                "cronograma acima são fixos e vinculantes, independendo "
                "do andamento da obra, salvo nas hipóteses de aditivo "
                "contratual devidamente assinado por ambas as partes.")

            p_f2 = doc.add_paragraph()
            p_f2.add_run("PARÁGRAFO SEGUNDO: ").bold = True
            p_f2.add_run(
                "O atraso no pagamento de qualquer parcela sujeitará "
                "o(a) CONTRATANTE à multa moratória de 2% (dois por "
                "cento) sobre o valor da parcela inadimplida, acrescida "
                "de juros de 1% (um por cento) ao mês, pro rata die, "
                "além de correção monetária pelo "
                f"{indice_correcao} desde a data do vencimento até "
                "a data do efetivo pagamento.")

            if tipo_pessoa_adm == 'PJ':
                p_nf = doc.add_paragraph()
                p_nf.add_run("PARÁGRAFO TERCEIRO — NOTA FISCAL: ").bold = True
                p_nf.add_run(
                    "O(A) CONTRATADO(A) deverá emitir Nota Fiscal de "
                    "Serviços (NFS-e) correspondente a cada parcela "
                    "paga, cujo comprovante de emissão deverá ser "
                    "enviado ao CONTRATANTE até o dia do vencimento.")

        elif metodo == "Eventos/Fases":
            # ── Eventos/Fases ────────────────────────────────────
            if entrada:
                valor_entrada_fmt = self._fmt_valor(entrada['valor'])
                intro_text = (
                    f"Como remuneração pelos serviços prestados, o(a) "
                    f"CONTRATANTE pagará ao(à) {denominacao_adm} o honorário "
                    f"total de {valor_fmt} ({valor_extenso}), sendo "
                    f"{valor_entrada_fmt} a título de sinal/entrada, pago na "
                    f"assinatura deste instrumento, e o restante distribuído "
                    f"em {len(parcelas_comuns)} evento(s)/fase(s) conforme "
                    "tabela abaixo, cada parcela devida após a conclusão e "
                    "ateste do respectivo marco de obra:")
            else:
                intro_text = (
                    f"Como remuneração pelos serviços prestados, o(a) "
                    f"CONTRATANTE pagará ao(à) {denominacao_adm} o honorário "
                    f"total de {valor_fmt} ({valor_extenso}), distribuído em "
                    f"{len(parcelas_comuns)} evento(s)/fase(s) conforme "
                    "tabela abaixo, sendo cada parcela devida após a "
                    "conclusão e ateste do respectivo marco de obra:")

            doc.add_paragraph(intro_text)

            todas_ev = []
            if entrada:
                todas_ev.append(entrada)
            todas_ev.extend(
                sorted(parcelas_comuns, key=lambda x: x['numero'] or 0)
            )

            if todas_ev:
                tem_datas = any(p.get('vencimento') for p in todas_ev)

                if tem_datas:
                    n_cols = 5
                    col_w  = [Inches(0.6), Inches(3.2), Inches(0.7), Inches(1.0), Inches(1.0)]
                    headers = ['Evento', 'Descrição', '%', 'Valor (R$)', 'Data Vencimento']
                else:
                    n_cols = 4
                    col_w  = [Inches(0.6), Inches(4.0), Inches(0.8), Inches(1.1)]
                    headers = ['Evento', 'Descrição', '%', 'Valor (R$)']

                tab = doc.add_table(rows=1, cols=n_cols)
                tab.style = 'Table Grid'

                for i, w in enumerate(col_w):
                    for cell in tab.columns[i].cells:
                        cell.width = w

                tbl = tab._tbl
                tblPr = tbl.tblPr
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '9360')
                tblW.set(qn('w:type'), 'dxa')
                tblPr.append(tblW)

                hdr = tab.rows[0].cells
                for i, h in enumerate(headers):
                    _cell_text(hdr[i], h, bold=True)

                total_tabela = 0.0
                for ev in todas_ev:
                    row_cells = tab.add_row().cells
                    for i, w in enumerate(col_w):
                        row_cells[i].width = w

                    if ev.get('eh_entrada') or ev['numero'] == 0:
                        num_display = 'ENTRADA'
                    else:
                        num_display = str(ev['numero'])

                    desc = str(ev.get('descricao') or '')
                    perc_val = ev['percentual']
                    if perc_val and perc_val < 1:
                        perc_val = perc_val * 100
                    perc_str = f"{perc_val:.1f}%" if perc_val else ''
                    val_fmt  = self._fmt_valor(ev['valor'])
                    dt_fmt   = self._fmt_data(ev.get('vencimento'))

                    _cell_text(row_cells[0], num_display)
                    _cell_text(row_cells[1], desc[:100] + ('…' if len(desc) > 100 else ''))
                    _cell_text(row_cells[2], perc_str)
                    _cell_text(row_cells[3], val_fmt)
                    if tem_datas:
                        _cell_text(row_cells[4], dt_fmt)

                    total_tabela += ev['valor']

                row_total = tab.add_row().cells
                for i, w in enumerate(col_w):
                    row_total[i].width = w
                _cell_text(row_total[0], 'TOTAL', bold=True)
                _cell_text(row_total[1], '',      bold=True)
                _cell_text(row_total[2], '',      bold=True)
                _cell_text(row_total[3], self._fmt_valor(total_tabela), bold=True)
                if tem_datas:
                    _cell_text(row_total[4], '', bold=True)

            par(space_before=8)

            p_e1 = doc.add_paragraph()
            p_e1.add_run("PARÁGRAFO PRIMEIRO: ").bold = True
            p_e1.add_run(
                "O pagamento de cada parcela ficará condicionado à "
                "conclusão e ao ateste do respectivo marco de obra, "
                "formalizado pelo(a) CONTRATANTE em até 5 (cinco) dias "
                "úteis após a comunicação de conclusão pelo(a) "
                f"{denominacao_adm}.")

            p_e2 = doc.add_paragraph()
            p_e2.add_run("PARÁGRAFO SEGUNDO: ").bold = True
            p_e2.add_run(
                "O não ateste sem justificativa técnica fundamentada "
                "no prazo acima importará em mora do(a) CONTRATANTE, "
                "sujeitando-o(a) à multa moratória de 0,5% (meio por "
                "cento) ao dia sobre o valor da parcela inadimplida, "
                "acrescida de correção monetária pelo "
                f"{indice_correcao}.")

            p_e3 = doc.add_paragraph()
            p_e3.add_run("PARÁGRAFO TERCEIRO: ").bold = True
            p_e3.add_run(
                "Na hipótese de paralisação da obra por mais de "
                "60 (sessenta) dias consecutivos por decisão unilateral "
                "do(a) CONTRATANTE, o(a) "
                f"{denominacao_adm} fará jus ao recebimento "
                "proporcional dos valores referentes às fases já "
                "concluídas, ainda que não formalmente atestadas.")

            if tipo_pessoa_adm == 'PJ':
                p_nf = doc.add_paragraph()
                p_nf.add_run("PARÁGRAFO QUARTO — NOTA FISCAL: ").bold = True
                p_nf.add_run(
                    "O(A) CONTRATADO(A) deverá emitir Nota Fiscal de "
                    "Serviços (NFS-e) para cada evento pago, "
                    "encaminhando-a ao(à) CONTRATANTE junto com a "
                    "solicitação de ateste.")

        # ── CLÁUSULA QUARTA — OBRIGAÇÕES DO CONTRATADO ───────────
        heading("CLÁUSULA QUARTA — OBRIGAÇÕES DO(A) CONTRATADO(A)")
        obrigacoes = [
            "executar os serviços com perícia técnica, observando "
            "as normas da ABNT e as boas práticas de engenharia;",
            "elaborar e enviar relatório quinzenal de prestação de "
            "contas ao(à) CONTRATANTE;",
            "manter sigilo sobre informações confidenciais da obra "
            "e do(a) CONTRATANTE;",
            "comunicar imediatamente ao(à) CONTRATANTE qualquer "
            "irregularidade ou risco identificado na obra;",
            "garantir o uso de EPIs pelos colaboradores em obra;",
            "não subcontratar total ou parcialmente a administração "
            "da obra sem autorização expressa e por escrito do(a) "
            "CONTRATANTE.",
        ]
        for letra, texto in zip('abcdef', obrigacoes):
            doc.add_paragraph(f"{letra}) {texto}")

        # ── CLÁUSULA QUINTA — OBRIGAÇÕES DO CONTRATANTE ──────────
        heading("CLÁUSULA QUINTA — OBRIGAÇÕES DO(A) CONTRATANTE")
        doc.add_paragraph(
            "a) fornecer ao(à) CONTRATADO(A) todos os projetos, "
            "aprovações e documentos necessários ao exercício da "
            "administração;")
        doc.add_paragraph("b) efetuar os pagamentos nos prazos estipulados;")
        doc.add_paragraph(
            "c) comunicar ao(à) CONTRATADO(A) qualquer alteração "
            "de escopo com antecedência mínima de 5 (cinco) dias úteis.")

        # ── CLÁUSULA SEXTA — RESPONSABILIDADE CIVIL ──────────────
        heading("CLÁUSULA SEXTA — RESPONSABILIDADE CIVIL")
        doc.add_paragraph(
            "O(A) CONTRATADO(A) responde por danos causados ao(à) "
            "CONTRATANTE ou a terceiros em decorrência de ação "
            "dolosa ou culposa sua ou de seus prepostos, devendo "
            "ressarcir integralmente os prejuízos comprovados.")

        # ── CLÁUSULA SÉTIMA — REAJUSTE E CORREÇÃO MONETÁRIA ──────
        heading("CLÁUSULA SÉTIMA — REAJUSTE E CORREÇÃO MONETÁRIA")

        if metodo == "Percentual da Quinzena":
            doc.add_paragraph(
                f"Por tratar-se de remuneração variável vinculada ao "
                "movimento financeiro da obra, não há reajuste anual "
                "automático do percentual contratado. Qualquer alteração "
                "do percentual somente poderá ser feita por Termo Aditivo "
                "assinado por ambas as partes.")
            p_cm1 = doc.add_paragraph()
            p_cm1.add_run("PARÁGRAFO ÚNICO: ").bold = True
            p_cm1.add_run(
                "Os valores em atraso serão corrigidos monetariamente "
                f"pelo {indice_correcao} (ou índice substituto legalmente "
                "adotado), desde a data do vencimento até o efetivo "
                "pagamento, acrescidos de multa de 2% e juros de 1% ao mês.")
        else:
            doc.add_paragraph(
                f"O valor global e as parcelas contratadas serão "
                f"reajustados anualmente, na data-base de aniversário "
                f"do contrato, pelo {indice_correcao} "
                "(Índice de correção contratado), acumulado nos "
                "12 (doze) meses imediatamente anteriores à data de "
                "reajuste, ou pelo índice que legalmente venha a "
                "substituí-lo.")
            p_cm1 = doc.add_paragraph()
            p_cm1.add_run("PARÁGRAFO PRIMEIRO: ").bold = True
            p_cm1.add_run(
                "O primeiro reajuste somente se aplicará após decorrido "
                "1 (um) ano da data de início da vigência deste "
                "instrumento.")
            p_cm2 = doc.add_paragraph()
            p_cm2.add_run("PARÁGRAFO SEGUNDO: ").bold = True
            p_cm2.add_run(
                "Os valores em atraso serão corrigidos monetariamente "
                f"pelo {indice_correcao}, acrescidos de multa de 2% "
                "(dois por cento) e juros de 1% (um por cento) ao mês, "
                "pro rata die, desde a data do vencimento até o "
                "efetivo pagamento.")
            p_cm3 = doc.add_paragraph()
            p_cm3.add_run("PARÁGRAFO TERCEIRO: ").bold = True
            p_cm3.add_run(
                "O reajuste será calculado e comunicado pelo(a) "
                f"{denominacao_adm} ao(à) CONTRATANTE com antecedência "
                "mínima de 30 (trinta) dias da data de sua aplicação, "
                "mediante apresentação do demonstrativo de cálculo.")

        # ── CLÁUSULA OITAVA — RESCISÃO ────────────────────────────
        heading("CLÁUSULA OITAVA — RESCISÃO E PENALIDADES")
        doc.add_paragraph(
            "Qualquer das partes poderá rescindir este contrato "
            "mediante notificação escrita com antecedência mínima "
            "de 30 (trinta) dias. A rescisão imotivada pelo(a) "
            "CONTRATANTE implicará no pagamento de todas as parcelas "
            "vencidas e das que venceriam nos 30 dias subsequentes. "
            "A rescisão por justa causa (descumprimento de cláusula) "
            "poderá ser imediata, sem esse ônus adicional.")

        p_mul = doc.add_paragraph()
        p_mul.add_run("PARÁGRAFO ÚNICO — MULTA: ").bold = True
        p_mul.add_run(
            f"O inadimplemento de qualquer cláusula sujeita a parte "
            f"infratora à multa não compensatória de {multa_fmt} "
            f"({multa_extenso}), sem prejuízo de perdas e danos.")

        # ── CLÁUSULA NONA — DISPOSIÇÕES GERAIS ───────────────────
        heading("CLÁUSULA NONA — DISPOSIÇÕES GERAIS")
        doc.add_paragraph(
            "a) Quaisquer alterações de valores ou escopo deverão "
            "ser objeto de Termo Aditivo;")
        doc.add_paragraph(
            "b) É vedado ao(à) CONTRATADO(A) utilizar trabalhadores "
            "menores de 18 anos;")
        doc.add_paragraph(
            "c) O(A) CONTRATANTE terá direito a ação regressiva "
            "em caso de condenação trabalhista oriunda de "
            "descumprimento pelo(a) CONTRATADO(A);")
        doc.add_paragraph(
            "d) Os comprovantes de transferência bancária constituem "
            "recibo de quitação dos valores pagos.")

        # ── CLÁUSULA DÉCIMA — FORO ────────────────────────────────
        heading("CLÁUSULA DÉCIMA — FORO")
        doc.add_paragraph(
            "As partes elegem o foro da Comarca de Belo Horizonte, "
            "Estado de Minas Gerais, para dirimir quaisquer "
            "controvérsias oriundas deste contrato, renunciando "
            "expressamente a qualquer outro, por mais privilegiado "
            "que seja.")

        # ── ENCERRAMENTO ─────────────────────────────────────────
        doc.add_paragraph(
            "E, por estarem assim justos e contratados, firmam o "
            "presente instrumento em duas (02) vias de igual teor "
            "e forma, na presença das testemunhas abaixo.")

        p_dt = doc.add_paragraph()
        p_dt.paragraph_format.space_before = Pt(10)
        p_dt.add_run(f"Belo Horizonte — MG, {data_extenso}.")

        # ── ASSINATURAS ───────────────────────────────────────────
        par(space_before=30)
        par("_" * 55)
        p_cn = doc.add_paragraph()
        p_cn.add_run(nome_cliente).bold = True

        par(space_before=30)
        par("_" * 55)
        p_adm2 = doc.add_paragraph()
        p_adm2.add_run(nome_adm).bold = True

        par(space_before=40)
        par("Testemunhas:", bold=True)
        par(space_before=15)
        par("_" * 55)
        par("RG n.º ")
        par(space_before=15)
        par("_" * 55)
        par("RG n.º ")

        # ── DADOS BANCÁRIOS ───────────────────────────────────────
        par(space_before=60)
        par("DADOS BANCÁRIOS PARA PAGAMENTO:", bold=True)
        p_adm3 = doc.add_paragraph()
        p_adm3.add_run(nome_adm).bold = True
        doc.add_paragraph(dados_banc or "Dados bancários não informados.")

        return doc

    # ── método público principal ─────────────────────────────────

    def gerar_contratos_do_contrato(
            self,
            nome_cliente: str,
            num_contrato: str,
            arquivo_cliente: Path,
            arquivo_clientes_geral: Path,
            arquivo_fornecedores: Path,
            pasta_saida: Path = None,
            metodo_override: str = None,
    ) -> list:
        """
        Lê o contrato num_contrato da planilha do cliente e gera
        um arquivo .docx para cada administrador cadastrado.

        Parâmetros
        ----------
        nome_cliente            : nome do cliente (para busca em Clientes.xlsx)
        num_contrato            : ex. '2024/08'  (sem sufixo J/F, busca ambos)
        arquivo_cliente         : caminho para CLIENTE.xlsx
        arquivo_clientes_geral  : caminho para Clientes.xlsx
        arquivo_fornecedores    : caminho para base_fornecedores.xlsx
        pasta_saida             : onde salvar os .docx (default: PASTA_CONTRATOS_ADM)
        metodo_override         : força um método ("Percentual da Quinzena" etc.)

        Retorna
        -------
        list[str] : caminhos dos arquivos gerados
        """
        if pasta_saida is None:
            if self.PASTA_CONTRATOS_ADM:
                pasta_saida = self.PASTA_CONTRATOS_ADM
            else:
                pasta_saida = arquivo_cliente.parent / "Contratos_ADM"
        pasta_saida = Path(pasta_saida)
        pasta_saida.mkdir(parents=True, exist_ok=True)

        # tentar carregar com sufixo J e F
        contratos_carregados = {}
        for sufixo in ('J', 'F', ''):
            num = f"{num_contrato}{sufixo}".strip()
            try:
                dados = self._ler_dados_contrato(arquivo_cliente, num)
                if dados['admins']:
                    contratos_carregados[num] = dados
            except Exception:
                pass
        if not contratos_carregados:
            dados = self._ler_dados_contrato(arquivo_cliente, num_contrato)
            if dados['admins']:
                contratos_carregados[num_contrato] = dados

        dados_cliente = self._ler_dados_cliente(arquivo_clientes_geral, nome_cliente)
        arquivos_gerados = []

        for num_c, contrato in contratos_carregados.items():
            for admin in contrato['admins']:
                dados_forn = self._ler_dados_fornecedor(
                    arquivo_fornecedores, admin['cnpj_cpf'])

                # determinar método
                if metodo_override:
                    metodo = metodo_override
                else:
                    metodo = self._detectar_metodo(contrato, admin)

                doc = self._build_doc(contrato, admin, dados_cliente,
                                      dados_forn, metodo)

                tp = admin['tipo_pessoa']
                num_c_safe = re.sub(r'[^\w]', '-', str(num_c))
                nome_seg   = re.sub(r'[^\w]', '_', admin['nome'])[:40]
                nome_arq   = (f"Contrato_ADM_{num_c_safe}_{tp}_{nome_seg}_"
                              f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
                caminho = pasta_saida / nome_arq
                doc.save(str(caminho))
                arquivos_gerados.append(str(caminho))

        return arquivos_gerados

    # ── detecção de método ───────────────────────────────────────

    def _detectar_metodo(self, contrato: dict, admin: dict) -> str:
        """
        Detecta o método de pagamento a partir dos dados da planilha,
        distinguindo claramente "Valor Fixo em Parcelas" de "Eventos/Fases".

        Regras:
        - Sem parcelas cadastradas + percentual > 0 e valor_total == 0
            → Percentual da Quinzena
        - Parcelas com DESCRICAO de evento/fase (texto não-genérico)
          E sem data de vencimento pré-definida
            → Eventos/Fases
        - Parcelas com DATA DE VENCIMENTO definida
            → Valor Fixo em Parcelas
        - num_parcelas > 1 sem outros critérios
            → Valor Fixo em Parcelas
        """
        cnpj_admin_limpo = re.sub(r'\D', '', str(admin['cnpj_cpf']))
        parcelas_admin = [
            p for p in contrato['parcelas']
            if re.sub(r'\D', '', str(p['cnpj_cpf'])) == cnpj_admin_limpo
        ]

        if not parcelas_admin:
            if admin['percentual'] > 0 and admin['valor_total'] == 0:
                return "Percentual da Quinzena"
            if admin['num_parcelas'] > 1:
                return "Valor Fixo em Parcelas"
            return "Eventos/Fases"

        # Parcelas com data de vencimento definida → Valor Fixo
        tem_data_venc = any(
            p['vencimento'] for p in parcelas_admin if not p.get('eh_entrada')
        )
        if tem_data_venc:
            return "Valor Fixo em Parcelas"

        # Parcelas com percentual e descrição não-genérica → Eventos/Fases
        tem_eventos = any(
            p['percentual'] and p['descricao'] and
            not re.match(r'^PARCELA\s+\d+$', p['descricao'].strip(), re.I)
            for p in parcelas_admin
            if not p.get('eh_entrada')
        )
        if tem_eventos:
            return "Eventos/Fases"

        # Fallback
        if admin['num_parcelas'] > 1:
            return "Valor Fixo em Parcelas"
        if admin['percentual'] > 0 and admin['valor_total'] == 0:
            return "Percentual da Quinzena"
        return "Eventos/Fases"

    @staticmethod
    def _extenso_percentual(valor: float) -> str:
        """Converte percentual para extenso. Ex: 7.0 → 'sete'"""
        try:
            from num2words import num2words
            inteiro = int(valor)
            decimal = round((valor - inteiro) * 100)
            if decimal == 0:
                return num2words(inteiro, lang='pt_BR')
            else:
                return (f"{num2words(inteiro, lang='pt_BR')} vírgula "
                        f"{num2words(decimal, lang='pt_BR')}")
        except Exception as e:
            import logging
            logging.getLogger("sistema").debug(
                f"num2words indisponível para percentual por extenso: "
                f"{type(e).__name__}: {e}"
            )
            return str(valor)
