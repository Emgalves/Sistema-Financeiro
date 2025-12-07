import os
import json
from datetime import datetime
from pathlib import Path
import pandas as pd
import numpy as np
from openpyxl import load_workbook

# Importações do python-docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

# Importar configurações do sistema
from src.config.config import (
    ARQUIVO_CLIENTES,
    ARQUIVO_FORNECEDORES,
    PASTA_CLIENTES,
    BASE_PATH
)

# Importar logger
from src.config.logger_config import system_logger, log_action
logger = system_logger.get_logger()

# Importar utils
from src.config.utils import (
    formatar_cnpj_cpf,
    buscar_dados_bancarios_fornecedor
)


class GeradorContrato:
    """Classe para geração de contratos em formato DOCX usando python-docx"""
    
    # Caminho do arquivo JSON de serviços
    SERVICOS_JSON_PATH = BASE_PATH / "servicos_construcao.json"
    
    # Pasta de contratos
    PASTA_CONTRATOS = PASTA_CLIENTES / "Contratos"
    
    def __init__(self):
        """Inicializa o gerador de contratos"""
        self.servicos_json = self._carregar_servicos()
        self._garantir_pasta_contratos()
        logger.info("GeradorContrato inicializado com sucesso (usando python-docx)")
    
    def _garantir_pasta_contratos(self):
        """Garante que a pasta de contratos existe"""
        try:
            self.PASTA_CONTRATOS.mkdir(parents=True, exist_ok=True)
            logger.info(f"Pasta de contratos verificada: {self.PASTA_CONTRATOS}")
        except Exception as e:
            logger.error(f"Erro ao criar pasta de contratos: {e}")
    
    def _carregar_servicos(self):
        """Carrega a lista de serviços do arquivo JSON"""
        json_path = self.SERVICOS_JSON_PATH
        
        logger.info(f"Carregando serviços de: {json_path}")
        
        # Se não existir, criar arquivo JSON básico
        if not json_path.exists():
            logger.warning("Arquivo de serviços não encontrado, criando padrão...")
            servicos_basicos = {
                "categorias": {
                    "fundacao": {
                        "nome": "Fundação e Estrutura",
                        "servicos": [
                            "escavação e movimento de terra",
                            "fundação em sapata corrida",
                            "fundação em radier",
                            "estrutura em concreto armado",
                            "estrutura metálica",
                            "laje pré-moldada",
                            "impermeabilização de fundação"
                        ]
                    },
                    "alvenaria": {
                        "nome": "Alvenaria e Vedação",
                        "servicos": [
                            "alvenaria de tijolo cerâmico",
                            "alvenaria de bloco de concreto",
                            "parede em drywall",
                            "chapisco e emboço interno",
                            "reboco fino interno"
                        ]
                    },
                    "cobertura": {
                        "nome": "Cobertura",
                        "servicos": [
                            "estrutura de madeira para telhado",
                            "cobertura em telha cerâmica",
                            "cobertura em telha metálica",
                            "calha e rufo",
                            "forro em PVC"
                        ]
                    },
                    "instalacoes": {
                        "nome": "Instalações",
                        "servicos": [
                            "instalação hidráulica completa",
                            "instalação elétrica completa",
                            "instalação de louças e metais"
                        ]
                    },
                    "revestimentos": {
                        "nome": "Revestimentos",
                        "servicos": [
                            "contrapiso",
                            "piso cerâmico",
                            "piso porcelanato",
                            "azulejo em parede",
                            "pintura em látex"
                        ]
                    }
                }
            }
            
            try:
                json_path.parent.mkdir(parents=True, exist_ok=True)
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(servicos_basicos, f, ensure_ascii=False, indent=2)
                logger.info("Arquivo de serviços criado com sucesso")
            except Exception as e:
                logger.error(f"Erro ao criar arquivo de serviços: {e}")
                return servicos_basicos
        
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                servicos = json.load(f)
            logger.info(f"Serviços carregados: {len(servicos.get('categorias', {}))} categorias")
            return servicos
        except Exception as e:
            logger.error(f"Erro ao ler arquivo de serviços: {e}")
            return {"categorias": {}}
    
    def listar_categorias_servicos(self):
        """Retorna lista de categorias de serviços disponíveis"""
        categorias = []
        for key, value in self.servicos_json.get('categorias', {}).items():
            categorias.append({
                'id': key,
                'nome': value['nome'],
                'qtd_servicos': len(value['servicos'])
            })
        return categorias
    
    def listar_servicos_categoria(self, categoria_id):
        """Retorna lista de serviços de uma categoria específica"""
        categorias = self.servicos_json.get('categorias', {})
        if categoria_id in categorias:
            return categorias[categoria_id]['servicos']
        return []
    
    def _get_safe_value(self, row, col_name, default=''):
        """Obtém valor de forma segura, tratando NaN e None"""
        try:
            value = row.get(col_name, default)
            # Verificar se é NaN (pandas/numpy)
            if pd.isna(value) or value is None:
                return default
            # Converter para string e limpar
            return str(value).strip()
        except:
            return default
    
    def formatar_cno(self, cno):
        """
        Formata CNO no padrão XX.XXX.XXXXX/XX
        CORRIGIDO: Verifica se JÁ está formatado antes de processar
        """
        try:
            # Se vazio, retornar vazio
            if not cno or pd.isna(cno):
                return ''
            
            # Converter para string
            cno_str = str(cno).strip()
            
            # Se já está formatado (contém pontos/barra), retornar como está
            if '.' in cno_str and '/' in cno_str:
                logger.info(f"CNO já formatado: {cno_str}")
                return cno_str
            
            # Remover caracteres não numéricos
            cno_limpo = ''.join(filter(str.isdigit, cno_str))
            
            # Se não tiver dígitos suficientes, retornar original
            if len(cno_limpo) < 12:
                logger.warning(f"CNO com poucos dígitos ({len(cno_limpo)}): {cno_str}")
                return cno_str
            
            # Garantir 13 dígitos (completar com zeros à esquerda se necessário)
            cno_limpo = cno_limpo.zfill(13)
            
            # Formatar: XX.XXX.XXXXX/XX
            cno_formatado = f"{cno_limpo[:2]}.{cno_limpo[2:5]}.{cno_limpo[5:10]}/{cno_limpo[10:12]}"
            logger.info(f"CNO formatado: {cno_str} → {cno_formatado}")
            return cno_formatado
        except Exception as e:
            logger.error(f"Erro ao formatar CNO: {e}")
            return str(cno) if cno else ''
    
    def obter_dados_cliente(self, nome_cliente):
        """Obtém dados do cliente da planilha - VERSÃO ROBUSTA"""
        try:
            logger.info(f"Obtendo dados do cliente: {nome_cliente}")
            
            # Ler planilha
            df = pd.read_excel(ARQUIVO_CLIENTES)
            
            # Verificar colunas disponíveis
            logger.info(f"Colunas disponíveis: {list(df.columns)}")
            
            # Tentar diferentes nomes para coluna de cliente
            col_cliente = None
            for nome_col in ['Nome', 'Cliente', 'cliente', 'CLIENTE', 'nome', 'NOME']:
                if nome_col in df.columns:
                    col_cliente = nome_col
                    break
            
            if not col_cliente:
                logger.error("Coluna de nome do cliente não encontrada")
                return None
            
            # Buscar cliente
            cliente_row = df[df[col_cliente] == nome_cliente]
            
            if cliente_row.empty:
                logger.error(f"Cliente '{nome_cliente}' não encontrado na planilha")
                return None
            
            cliente = cliente_row.iloc[0]
            
            # Obter dados com tratamento robusto
            # IMPORTANTE: Planilha de clientes tem coluna 'CPF', não 'CPF/CNPJ'
            # NÃO usar _get_safe_value para CPF porque ele converte para string antes
            cpf_raw = cliente['CPF'] if 'CPF' in cliente.index else None
            
            logger.info(f"CPF bruto obtido: {cpf_raw} (tipo: {type(cpf_raw)})")
            
            # Se CPF veio como número (float), converter para string SEM .0
            if cpf_raw is not None and not pd.isna(cpf_raw):
                # Converter para string removendo .0 se for float
                if isinstance(cpf_raw, float):
                    cpf_str = str(int(cpf_raw))  # int() remove o .0
                    logger.info(f"CPF convertido de float: {cpf_raw} → {cpf_str}")
                elif isinstance(cpf_raw, int):
                    cpf_str = str(cpf_raw)
                    logger.info(f"CPF convertido de int: {cpf_raw} → {cpf_str}")
                else:
                    cpf_str = str(cpf_raw).strip()
                    logger.info(f"CPF como string: {cpf_str}")
                
                # Remover qualquer caractere não numérico (caso tenha pontos ou traços)
                apenas_numeros = ''.join(filter(str.isdigit, cpf_str))
                logger.info(f"CPF apenas números: {apenas_numeros}")
                
                # Formatar manualmente
                if len(apenas_numeros) == 11:  # CPF
                    cpf_formatado = f"{apenas_numeros[:3]}.{apenas_numeros[3:6]}.{apenas_numeros[6:9]}-{apenas_numeros[9:11]}"
                    logger.info(f"✅ CPF formatado: {cpf_formatado}")
                elif len(apenas_numeros) == 14:  # CNPJ
                    cpf_formatado = f"{apenas_numeros[:2]}.{apenas_numeros[2:5]}.{apenas_numeros[5:8]}/{apenas_numeros[8:12]}-{apenas_numeros[12:14]}"
                    logger.info(f"✅ CNPJ formatado: {cpf_formatado}")
                else:
                    logger.warning(f"⚠️ CPF/CNPJ com tamanho inválido: {len(apenas_numeros)} dígitos")
                    cpf_formatado = cpf_str
            else:
                logger.warning("⚠️ CPF não informado ou vazio")
                cpf_formatado = 'não informado'
            
            dados = {
                'nome': self._get_safe_value(cliente, col_cliente, nome_cliente),
                'cnpj_cpf': cpf_formatado,
                'cno': self.formatar_cno(self._get_safe_value(cliente, 'CNO', '')),
                'estado_civil': self._get_safe_value(cliente, 'Estado Civil', 'não informado'),
                'endereco': self._get_safe_value(cliente, 'Endereço', 'não informado'),
                'cidade': self._get_safe_value(cliente, 'Cidade', 'Belo Horizonte'),
                'estado': self._get_safe_value(cliente, 'Estado', 'MG')
            }
            
            logger.info(f"✅ Dados do cliente processados - CPF final: {dados['cnpj_cpf']}")
            
            logger.info(f"Dados do cliente obtidos com sucesso: {dados['nome']}")
            return dados
            
        except Exception as e:
            logger.error(f"Erro ao obter dados do cliente: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def obter_dados_fornecedor(self, cnpj_cpf):
        """
        Obtém dados do fornecedor da planilha pelo CPF/CNPJ
        CORRIGIDO: Busca correta do endereço na coluna 'ENDEREÇO'
        """
        try:
            logger.info(f"Obtendo dados do fornecedor: {cnpj_cpf}")
            
            df = pd.read_excel(ARQUIVO_FORNECEDORES)
            
            # Buscar fornecedor
            fornecedor_row = df[df['CNPJ/CPF'].astype(str).str.replace(r'\D', '', regex=True) == 
                              cnpj_cpf.replace('.', '').replace('-', '').replace('/', '')]
            
            if fornecedor_row.empty:
                logger.error(f"Fornecedor '{cnpj_cpf}' não encontrado na planilha")
                return None
            
            fornecedor = fornecedor_row.iloc[0]
            
            # CORREÇÃO: Buscar endereço na coluna correta 'ENDEREÇO' (índice 15)
            endereco_raw = fornecedor.get('ENDEREÇO', None)
            
            # Tratar endereço vazio/NaN
            if pd.isna(endereco_raw) or not str(endereco_raw).strip():
                endereco = 'não informado'
                logger.warning(f"Endereço não encontrado para fornecedor {cnpj_cpf}")
            else:
                endereco = str(endereco_raw).strip()
                logger.info(f"✅ Endereço encontrado: {endereco}")
            
            dados = {
                'nome': self._get_safe_value(fornecedor, 'NOME', 'não informado'),
                'cnpj_cpf': formatar_cnpj_cpf(cnpj_cpf),
                'endereco': endereco
            }
            
            logger.info(f"Dados do fornecedor obtidos: {dados['nome']}")
            return dados
            
        except Exception as e:
            logger.error(f"Erro ao obter dados do fornecedor: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def obter_dados_fornecedor_por_nome(self, nome_fornecedor):
        """
        Obtém dados do fornecedor da planilha pelo nome
        CORRIGIDO: Busca correta do endereço
        """
        try:
            logger.info(f"Obtendo dados do fornecedor por nome: {nome_fornecedor}")
            
            df = pd.read_excel(ARQUIVO_FORNECEDORES)
            logger.info(f"Colunas disponíveis na planilha de fornecedores: {list(df.columns)}")
            
            # Colunas reais da planilha
            col_nome = 'NOME'  # Nome exato da coluna
            col_cnpj = 'CNPJ/CPF'  # Nome exato da coluna
            
            if col_nome not in df.columns:
                logger.error(f"Coluna '{col_nome}' não encontrada na planilha")
                logger.error(f"Colunas disponíveis: {list(df.columns)}")
                return None
            
            # Buscar fornecedor pelo nome (case insensitive)
            fornecedor_row = df[df[col_nome].astype(str).str.strip().str.upper() == nome_fornecedor.strip().upper()]
            
            if fornecedor_row.empty:
                logger.error(f"Fornecedor '{nome_fornecedor}' não encontrado na planilha")
                logger.error(f"Primeiros 5 fornecedores: {df[col_nome].head().tolist()}")
                return None
            
            fornecedor = fornecedor_row.iloc[0]
            logger.info(f"✅ Fornecedor encontrado na planilha")
            
            # Obter CPF/CNPJ (pode vir como número inteiro)
            cnpj_cpf_raw = fornecedor[col_cnpj] if col_cnpj in fornecedor.index else None
            
            if pd.isna(cnpj_cpf_raw):
                logger.error(f"CPF/CNPJ está vazio para fornecedor '{nome_fornecedor}'!")
                cnpj_cpf_formatado = 'não informado'
            else:
                # Converter para string (pode estar como int)
                cnpj_cpf_str = str(int(cnpj_cpf_raw)) if isinstance(cnpj_cpf_raw, (int, float)) else str(cnpj_cpf_raw)
                
                # Tentar formatar usando a função do utils
                cnpj_cpf_formatado = formatar_cnpj_cpf(cnpj_cpf_str)
                
                # Se a função não formatou (retornou igual), formatar manualmente
                if cnpj_cpf_formatado == cnpj_cpf_str:
                    # Remover caracteres não numéricos
                    apenas_numeros = ''.join(filter(str.isdigit, cnpj_cpf_str))
                    
                    if len(apenas_numeros) == 14:  # CNPJ
                        # Formato: XX.XXX.XXX/XXXX-XX
                        cnpj_cpf_formatado = f"{apenas_numeros[:2]}.{apenas_numeros[2:5]}.{apenas_numeros[5:8]}/{apenas_numeros[8:12]}-{apenas_numeros[12:14]}"
                    elif len(apenas_numeros) == 11:  # CPF
                        # Formato: XXX.XXX.XXX-XX
                        cnpj_cpf_formatado = f"{apenas_numeros[:3]}.{apenas_numeros[3:6]}.{apenas_numeros[6:9]}-{apenas_numeros[9:11]}"
                    else:
                        # Se não for CNPJ nem CPF, manter como está
                        cnpj_cpf_formatado = cnpj_cpf_str
                
                logger.info(f"CPF/CNPJ: {cnpj_cpf_raw} → {cnpj_cpf_formatado}")
            
            # CORREÇÃO: Obter endereço da coluna correta
            endereco_raw = fornecedor.get('ENDEREÇO', None)
            
            if pd.isna(endereco_raw) or not str(endereco_raw).strip():
                endereco = 'não informado'
                logger.warning(f"Endereço não encontrado para fornecedor {nome_fornecedor}")
            else:
                endereco = str(endereco_raw).strip()
                logger.info(f"✅ Endereço encontrado: {endereco}")
            
            # Obter outros dados
            razao_social = fornecedor.get('RAZÃO SOCIAL', nome_fornecedor)
            dados_bancarios = fornecedor.get('DADOS BANCÁRIOS', '')
            
            dados = {
                'nome': str(fornecedor[col_nome]).strip(),
                'razao_social': str(razao_social).strip() if not pd.isna(razao_social) else nome_fornecedor,
                'cnpj_cpf': cnpj_cpf_formatado,
                'endereco': endereco,
                'dados_bancarios': str(dados_bancarios).strip() if not pd.isna(dados_bancarios) else 'não informado'
            }
            
            logger.info(f"✅ Dados obtidos: Nome={dados['nome']}, CNPJ/CPF={dados['cnpj_cpf']}, Endereço={dados['endereco']}")
            return dados
            
        except Exception as e:
            logger.error(f"❌ Erro ao obter dados do fornecedor por nome: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def formatar_data_extenso(self, data_str):
        """Converte data de DD/MM/YYYY para extenso"""
        try:
            meses = {
                1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
                5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
                9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
            }
            
            data = datetime.strptime(data_str, '%d/%m/%Y')
            return f"{data.day} de {meses[data.month]} de {data.year}"
        except:
            return data_str
    
    def numero_por_extenso(self, valor):
        """
        Converte número para extenso (simplificado)
        CORRIGIDO: Não multiplica mais por 10
        """
        try:
            from num2words import num2words
            # CORREÇÃO: valor já vem como float correto, não precisa ajustar
            valor_float = float(valor)
            return num2words(valor_float, lang='pt_BR', to='currency')
        except:
            return f"{valor} reais"
    
    def concatenar_servicos(self, lista_servicos):
        """Concatena lista de serviços de forma gramaticalmente correta"""
        if not lista_servicos:
            return ""
        
        if len(lista_servicos) == 1:
            return lista_servicos[0]
        
        if len(lista_servicos) == 2:
            return f"{lista_servicos[0]} e {lista_servicos[1]}"
        
        # Múltiplos serviços: "x, y, z e w"
        servicos_texto = ", ".join(lista_servicos[:-1])
        return f"{servicos_texto} e {lista_servicos[-1]}"
    
    def _configurar_estilos(self, doc):
        """
        Configura os estilos do documento
        CORRIGIDO: CLÁUSULA em preto (não azul) e maior que PARÁGRAFO
        """
        # Estilo Normal
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Arial'
        font_normal.size = Pt(11)
        
        paragraph_format = style_normal.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)
        
        # Estilo Título
        try:
            style_title = doc.styles['Title']
        except KeyError:
            style_title = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        
        font_title = style_title.font
        font_title.name = 'Arial'
        font_title.size = Pt(14)
        font_title.bold = True
        
        paragraph_format_title = style_title.paragraph_format
        paragraph_format_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format_title.space_before = Pt(12)
        paragraph_format_title.space_after = Pt(12)
        
        # CORREÇÃO: Estilo Heading 1 - CLÁUSULA (maior e em preto)
        style_h1 = doc.styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.name = 'Arial'
        font_h1.size = Pt(12)  # Maior que parágrafo
        font_h1.bold = True
        font_h1.color.rgb = RGBColor(0, 0, 0)  # PRETO, não azul
        
        paragraph_format_h1 = style_h1.paragraph_format
        paragraph_format_h1.space_before = Pt(12)  # Mais espaço
        paragraph_format_h1.space_after = Pt(6)
    
    def gerar_contrato(self, nome_cliente_ou_dict=None, cnpj_fornecedor=None, descricao_servicos=None, 
                      data_inicio=None, data_fim=None, valor_global=None, prazo_dias=None, multa_valor=None, observacoes=''):
        """
        Gera contrato em formato DOCX usando python-docx
        
        CORREÇÕES IMPLEMENTADAS:
        1. CNO: verifica se já está formatado
        2. Endereço fornecedor: busca na coluna 'ENDEREÇO' corretamente
        3. Data do Contrato: usa data_inicio ao invés de data atual
        4. Prazo: usa valor informado no formulário (prazo_dias)
        5. Valor Global por extenso: corrigido (não multiplica por 10)
        6. Multa: usa valor informado no formulário
        7. CLÁUSULA: em preto e visualmente maior que PARÁGRAFO
        8. Linhas de assinatura: sem quebra, com espaçamento adequado
        9. Linha testemunhas: com separação
        10. Espaçamento antes de DADOS BANCÁRIOS
        
        Aceita dois formatos:
        1. Dicionário completo (formato novo)
        2. Parâmetros individuais (compatibilidade)
        
        Args:
            nome_cliente_ou_dict: Nome do cliente OU dicionário com todos os dados
            cnpj_fornecedor: CPF/CNPJ do fornecedor (ignorado se dict)
            descricao_servicos: Descrição dos serviços (ignorado se dict)
            data_inicio: Data de início (ignorado se dict)
            data_fim: Data de fim (ignorado se dict)
            valor_global: Valor global do contrato (ignorado se dict)
            prazo_dias: Prazo em dias úteis (NOVO - usa valor do formulário)
            multa_valor: Valor da multa (NOVO - usa valor do formulário)
            observacoes: Observações adicionais (ignorado se dict)
            
        Returns:
            str: Caminho do arquivo gerado ou None se erro
        """
        try:
            logger.info("=== INICIANDO GERAÇÃO DE CONTRATO ===")
            
            # Verificar se é dicionário (novo formato) ou parâmetros individuais
            if isinstance(nome_cliente_ou_dict, dict):
                # Formato novo: dict completo
                dados_contrato = nome_cliente_ou_dict
                
                logger.info(f"📋 Dados do contrato recebidos: {list(dados_contrato.keys())}")
                
                # Extrair dados necessários
                nome_cliente = dados_contrato.get('cliente_nome')
                cnpj_fornecedor = dados_contrato.get('fornecedor_cnpj_cpf')
                descricao_servicos = dados_contrato.get('descricao', '')
                data_inicio = dados_contrato.get('data_inicio')
                data_fim = dados_contrato.get('data_fim')
                valor_global = dados_contrato.get('valor')
                
                # CORREÇÃO: Usar prazo e multa do formulário
                # Aceitar tanto 'prazo_dias' quanto 'dias' para compatibilidade
                prazo_dias = dados_contrato.get('prazo_dias') or dados_contrato.get('dias')
                multa_valor = dados_contrato.get('multa', None)
                
                logger.info(f"Usando formato dict - Cliente: {nome_cliente}")
                logger.info(f"Fornecedor CPF/CNPJ do dict: {cnpj_fornecedor}")
                logger.info(f"Prazo informado: {prazo_dias} dias")
                logger.info(f"Multa informada: {multa_valor}")
                
                # Validar campos obrigatórios
                if not nome_cliente:
                    logger.error("Campo 'cliente_nome' está vazio!")
                    raise ValueError("Nome do cliente não informado")
                
                if not dados_contrato.get('fornecedor_nome'):
                    logger.error("Campo 'fornecedor_nome' está vazio!")
                    raise ValueError("Nome do fornecedor não informado")
                
                if not cnpj_fornecedor:
                    logger.warning("Campo 'fornecedor_cnpj_cpf' está vazio! Usando 'não informado'")
                    cnpj_fornecedor = 'não informado'
                
                # Dados já vêm formatados no dict
                dados_cliente = {
                    'nome': dados_contrato.get('cliente_nome'),
                    'cnpj_cpf': dados_contrato.get('cliente_cpf', dados_contrato.get('cliente_cnpj_cpf', '')),
                    'cno': dados_contrato.get('cliente_cno', ''),
                    'estado_civil': dados_contrato.get('cliente_estado_civil', 'não informado'),
                    'endereco': dados_contrato.get('cliente_endereco', 'não informado'),
                    'cidade': dados_contrato.get('cidade', 'Belo Horizonte'),
                    'estado': dados_contrato.get('estado', 'MG')
                }
                
                dados_fornecedor = {
                    'nome': dados_contrato.get('fornecedor_nome', 'não informado'),
                    'cnpj_cpf': cnpj_fornecedor,
                    'endereco': dados_contrato.get('fornecedor_endereco', 'não informado')
                }
                
                logger.info(f"✅ Dados cliente processados: {dados_cliente.get('nome')}")
                logger.info(f"✅ Dados fornecedor processados: {dados_fornecedor.get('nome')} - {dados_fornecedor.get('cnpj_cpf')}")
                
                dados_bancarios = dados_contrato.get('dados_bancarios', 'Dados bancários não cadastrados')
                endereco_obra = dados_contrato.get('endereco_obra', dados_cliente['endereco'])
                
            else:
                # Formato antigo: parâmetros individuais
                nome_cliente = nome_cliente_ou_dict
                
                logger.info(f"Usando formato legado - Cliente: {nome_cliente}")
                
                # Obter dados do cliente
                dados_cliente = self.obter_dados_cliente(nome_cliente)
                if not dados_cliente:
                    logger.error("Não foi possível obter dados do cliente")
                    return None
                
                # Obter dados do fornecedor
                dados_fornecedor = self.obter_dados_fornecedor(cnpj_fornecedor)
                if not dados_fornecedor:
                    logger.error("Não foi possível obter dados do fornecedor")
                    return None
                
                # Buscar dados bancários
                dados_bancarios = buscar_dados_bancarios_fornecedor(cnpj_fornecedor)
                if not dados_bancarios:
                    dados_bancarios = "Dados bancários não cadastrados"
                
                endereco_obra = dados_cliente['endereco']
            
            # CORREÇÃO: Usar prazo informado no formulário (dias ÚTEIS)
            # Converter para int se vier como string
            if prazo_dias is not None and str(prazo_dias).strip() and str(prazo_dias).strip() != '0':
                try:
                    dias = int(prazo_dias)
                    logger.info(f"✅ Usando prazo do formulário: {dias} dias úteis")
                except (ValueError, TypeError):
                    logger.warning(f"⚠️ Prazo inválido: '{prazo_dias}', calculando dias úteis")
                    prazo_dias = None  # Forçar cálculo
            
            if prazo_dias is None or str(prazo_dias).strip() == '0':
                # Fallback: calcular dias ÚTEIS (não corridos) se prazo não foi informado
                try:
                    from datetime import datetime as dt
                    import numpy as np
                    
                    dt_inicio = datetime.strptime(data_inicio, '%d/%m/%Y')
                    dt_fim = datetime.strptime(data_fim, '%d/%m/%Y')
                    
                    # Calcular dias ÚTEIS usando numpy.busday_count
                    dias = np.busday_count(dt_inicio.date(), dt_fim.date())
                    logger.warning(f"⚠️ Prazo não informado, calculando dias ÚTEIS: {dias}")
                except Exception as e:
                    logger.error(f"Erro ao calcular dias úteis: {e}")
                    # Fallback final: calcular aproximadamente (dias corridos * 0.71)
                    try:
                        dt_inicio = datetime.strptime(data_inicio, '%d/%m/%Y')
                        dt_fim = datetime.strptime(data_fim, '%d/%m/%Y')
                        dias_corridos = (dt_fim - dt_inicio).days
                        dias = int(dias_corridos * 0.71)  # Aproximação: 5 dias úteis em 7 corridos
                        logger.warning(f"⚠️ Usando aproximação de dias úteis: {dias} (baseado em {dias_corridos} corridos)")
                    except:
                        dias = 30
                        logger.warning(f"⚠️ Erro ao calcular prazo, usando padrão: {dias}")
            
            # Preparar valores
            valor_limpo = str(valor_global).replace('R$', '').replace('.', '').replace(',', '.').strip()
            valor_float = float(valor_limpo)
            valor_formatado = f"R$ {valor_float:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            
            # CORREÇÃO: valor por extenso correto (não multiplica por 10)
            valor_extenso = self.numero_por_extenso(valor_float)
            
            # CORREÇÃO: Usar multa informada no formulário
            if multa_valor is not None:
                multa_limpo = str(multa_valor).replace('R$', '').replace('.', '').replace(',', '.').strip()
                multa_float = float(multa_limpo)
                logger.info(f"✅ Usando multa do formulário: R$ {multa_float}")
            else:
                # Fallback: calcular 10% se multa não foi informada
                multa_float = valor_float * 0.10
                logger.warning(f"⚠️ Multa não informada, calculando 10%: R$ {multa_float}")
            
            multa_formatada = f"R$ {multa_float:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            
            # CORREÇÃO: multa por extenso correto (não multiplica por 10)
            multa_extenso = self.numero_por_extenso(multa_float)
            
            # CORREÇÃO: Data por extenso da data de INÍCIO (não data atual)
            data_extenso = self.formatar_data_extenso(data_inicio)
            logger.info(f"✅ Data do contrato: {data_inicio} ({data_extenso})")
            
            # Criar documento
            doc = Document()
            self._configurar_estilos(doc)
            
            # Configurar margens (1 polegada = 1440 twips)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # TÍTULO
            titulo = doc.add_paragraph()
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_titulo = titulo.add_run("CONTRATO PARTICULAR DE PRESTAÇÃO DE SERVIÇOS POR EMPREITADA")
            run_titulo.bold = True
            run_titulo.font.size = Pt(14)
            
            # INTRODUÇÃO
            p1 = doc.add_paragraph()
            p1.add_run(f"Aos {data_inicio}, nesta cidade de {dados_cliente['cidade']}, entre partes, de um lado: ")
            
            # CONTRATANTE
            p2 = doc.add_paragraph()
            run_cliente = p2.add_run(dados_cliente['nome'])
            run_cliente.bold = True
            p2.add_run(f", pessoa física devidamente inscrita sob o CNO n.º {dados_cliente['cno']} e CPF nº {dados_cliente['cnpj_cpf']}, {dados_cliente['estado_civil']}, residente na {dados_cliente['endereco']}, doravante denominada CONTRATANTE e, de outro ")
            
            # CONTRATADA
            run_fornecedor = p2.add_run(dados_fornecedor['nome'])
            run_fornecedor.bold = True
            p2.add_run(f", pessoa física devidamente inscrita sob o CPF n.º {dados_fornecedor['cnpj_cpf']} com residência na {dados_fornecedor['endereco']}, doravante denominado simplesmente de CONTRATADA, ambas representadas por seus representantes legais que ao final firmam o presente contrato, tem entre si, justo e contratado o presente, que se regerá pelas seguintes Cláusulas e Condições:")
            
            # CLÁUSULA PRIMEIRA - OBJETO
            doc.add_heading('CLÁUSULA PRIMEIRA - OBJETO', level=1)
            doc.add_paragraph(f"O presente contrato tem como OBJETO a prestação de serviços especializados em {descricao_servicos}, bem como todos os trabalhos e atividades necessárias para sua conclusão.")
            
            p_par1 = doc.add_paragraph()
            run_par1 = p_par1.add_run("PARÁGRAFO PRIMEIRO: ")
            run_par1.bold = True
            run_par1.font.size = Pt(11)  # Menor que CLÁUSULA
            
            doc.add_paragraph(f"Os serviços deverão ser prestados no imóvel situado à {endereco_obra}")
            
            p_par2 = doc.add_paragraph()
            run_par2 = p_par2.add_run("PARÁGRAFO SEGUNDO: ")
            run_par2.bold = True
            run_par2.font.size = Pt(11)
            
            doc.add_paragraph("A contratada prestará os serviços constantes em orçamento e/ou descritivo de atividades na modalidade por empreitada de forma autônoma, sem qualquer exclusividade, podendo desempenhar atividades para terceiros em geral, simultaneamente ou não.")
            
            # CLÁUSULA SEGUNDA - SERVIÇOS
            doc.add_heading('CLÁUSULA SEGUNDA - SERVIÇOS', level=1)
            doc.add_paragraph("Os serviços acima mencionados serão prestados pela contratada através de seus prepostos ou empregados devidamente registrados, sem qualquer vinculação com a contratante.")
            
            p_par_primeiro = doc.add_paragraph()
            run_par_primeiro = p_par_primeiro.add_run("PARÁGRAFO PRIMEIRO:")
            run_par_primeiro.bold = True
            run_par_primeiro.font.size = Pt(11)
            
            doc.add_paragraph("O Contratado obrigar-se-á:")
            
            p_a = doc.add_paragraph()
            p_a.add_run("a) executar os serviços autônomos com toda a perfeição técnica na forma e modo ajustados, dentro das normas e especificações técnicas aplicáveis à espécie e ")
            run_underline = p_a.add_run("em estrito cumprimento dos detalhes, projetos e especificações, dando plena e total garantia dos mesmos;")
            run_underline.underline = True
            
            doc.add_paragraph("b) fornecer toda mão-de-obra necessária à execução e entrega dos serviços no prazo estabelecido, devendo registrar todos os trabalhadores em seu nome, obrigando-se pelos salários dos empregados que o mesmo utilizar na obra, comprometendo-se a respeitar as normas trabalhistas, de segurança do trabalho e previdenciárias vigentes;")
            
            doc.add_paragraph("c) fornecer todas as ferramentas necessárias para a execução dos serviços contratados;")
            
            doc.add_paragraph("d) corrigir, por sua conta e risco, qualquer defeito constatado durante a construção ou instalação ou execução e/ou oriundo de imperfeição de serviços;")
            
            doc.add_paragraph("e) pagamento dos encargos sociais, previdenciários e trabalhistas dos colaboradores utilizados na execução dos serviços ora contratados;")
            
            doc.add_paragraph("f) garantir a solidez e estabilidade do serviço prestado, assumindo, por ela, inteira responsabilidade, pelos danos oriundos de sua negligência, imprudência ou imperícia nos termos do Código Civil Brasileiro;")
            
            doc.add_paragraph("g) manter, por sua conta, seguro contra acidentes de trabalho em nome de todos os colaboradores que trabalharem na obra;")
            
            doc.add_paragraph("h) Fornecer, zelar e garantir o uso de equipamentos de proteção individuais e coletivos na execução dos serviços e ambiente da obra, como forma de atender todas as normas de segurança e higiene do trabalho vigentes e pertinentes ao ramo de sua atividade.")
            
            doc.add_paragraph("i) Avaliar e mitigar os riscos para iniciar a execução dos trabalhos sendo que na possibilidade de verificar o menor risco de acidente deverá comunicar o contratante sem adentrar ao ambiente de prestação de serviços, medida necessária para garantir segurança aos seus colaboradores.")
            
            p_par_segundo = doc.add_paragraph()
            run_par_segundo = p_par_segundo.add_run("PARÁGRAFO SEGUNDO:")
            run_par_segundo.bold = True
            run_par_segundo.font.size = Pt(11)
            
            doc.add_paragraph("São obrigações exclusivas do contratante:")
            doc.add_paragraph("a) Fornecer todos os detalhes, projetos e especificações para a perfeita execução dos serviços;")
            doc.add_paragraph("b) Efetuar o pagamento na forma e modo aprazados.")
            
            # CLÁUSULA TERCEIRA - PRAZO
            doc.add_heading('CLÁUSULA TERCEIRA - PRAZO', level=1)
            
            p_prazo = doc.add_paragraph()
            p_prazo.add_run("Os serviços ora contratados serão executados/prestados até o limite de ")
            run_dias = p_prazo.add_run(f"{dias} dias")
            run_dias.bold = True
            p_prazo.add_run(", iniciando-se a contagem com a assinatura deste.")
            
            doc.add_paragraph(f"Iniciando-se a contagem com a entrada no campo de obras que está prevista para {data_inicio} e encerrando-se em {data_fim}.")
            
            # CLÁUSULA QUARTA - REMUNERAÇÃO
            doc.add_heading('CLÁUSULA QUARTA -- REMUNERAÇÃO', level=1)
            doc.add_paragraph(f"Como remuneração pelos serviços a serem prestados, os contratantes pagarão ao contratado, mediante depósito/transferência bancária, o valor de {valor_formatado} ({valor_extenso}), para pagamento integral dos serviços contratados por este instrumento valores fixos e irreajustáveis, valores que serão pagos mediante medição, após sua execução. Os valores convencionados deverão ser pagos na medida e prazos em que a prestação de serviços se desenvolver, podendo o contratante reter o pagamento, sem nenhum ônus, caso o serviço não seja prestado adequadamente ou integralmente nos moldes e diretrizes estabelecidas pelas partes e projetos de conhecimento.")
            
            p_par_prim = doc.add_paragraph()
            run_par_prim = p_par_prim.add_run("PARÁGRAFO PRIMEIRO")
            run_par_prim.bold = True
            run_par_prim.font.size = Pt(11)
            
            doc.add_paragraph("A remuneração pelos serviços contratados inclui todos os encargos trabalhistas, sociais, previdenciários, securitários e outros não nominados, gastos e despesas relativos ao exercício dos serviços contratados, por mais especiais que sejam, nada mais sendo devido pelo contratante ao contratado, a qualquer título.")
            
            p_par_seg = doc.add_paragraph()
            run_par_seg = p_par_seg.add_run("PARÁGRAFO SEGUNDO")
            run_par_seg.bold = True
            run_par_seg.font.size = Pt(11)
            
            doc.add_paragraph("O presente contrato não implica em qualquer vínculo empregatício do contratado, de seus prepostos ou colaboradores pelos serviços prestados ao contratante.")
            
            p_par_terc = doc.add_paragraph()
            run_par_terc = p_par_terc.add_run("PARÁGRAFO TERCEIRO")
            run_par_terc.bold = True
            run_par_terc.font.size = Pt(11)
            
            doc.add_paragraph("Os comprovantes de transferência servirão como recibo de quitação dos valores eventualmente pagos à Contratada.")
            
            # CLÁUSULA QUINTA - DISPOSIÇÕES GERAIS
            doc.add_heading('CLÁUSULA QUINTA - DISPOSIÇÕES GERAIS', level=1)
            doc.add_paragraph("a) As alterações de valores que venham a ser discutidos e aprovados pelas partes, deverão necessariamente ser objeto de Termo Aditivo.")
            doc.add_paragraph("b) A transferência ou cessão dos serviços de que trata o presente instrumento depende do consentimento expresso deste contratante, bem como a aditivo contratual, constando assinatura do contratante.")
            doc.add_paragraph("c) É expressamente vedada à Contratada a utilização de trabalhadores menores, púberes ou impúberes, para a prestação dos serviços.")
            doc.add_paragraph("d) Ao contratante fica ressalvado o direito à ação regressiva em face do contratado e ainda, a retenção da importância devida, em razão da quitação de eventuais obrigações trabalhistas dos empregados do contratado que eventualmente venha a sofrer em decorrência de acordos ou decisões judiciais.")
            doc.add_paragraph("e) Fica assegurado o direito do contratante ao ressarcimento dos danos sofridos em virtude de interpelação judicial em razão de obrigação não cumprida pelo contratado, inclusive eventuais despesas com honorários advocatícios contratuais.")
            
            # CLÁUSULA SEXTA - DOS PREJUÍZOS
            doc.add_heading('CLÁUSULA SEXTA -- DOS PREJUÍZOS', level=1)
            doc.add_paragraph("A contratada responderá por qualquer prejuízo que direta ou indiretamente cause ao contratante ou a terceiros, seja por ação ou omissão, sua ou de seus prepostos, empregados ou colaboradores.")
            
            # CLÁUSULA SÉTIMA - DA RESCISÃO
            doc.add_heading('CLÁUSULA SÉTIMA -- DA RESCISÃO', level=1)
            doc.add_paragraph("Serão casos de rescisão contratual:")
            doc.add_paragraph("a) a desistência de uma das partes antes de iniciada a prestação de serviços;")
            doc.add_paragraph("b) a falha do Contratado em executar os trabalhos ora especificados, nas condições estipuladas ou paralisação da obra por mais de 7 (sete) dias sem relevante razão;")
            doc.add_paragraph("c) qualquer outro fato ou ato que, por culpa ou dolo de uma das partes, impossibilite a execução do presente contrato.")
            
            p_par_unico = doc.add_paragraph()
            run_par_unico = p_par_unico.add_run("PARÁGRAFO ÚNICO -- ")
            run_par_unico.bold = True
            run_par_unico.font.size = Pt(11)
            p_par_unico.add_run(f"Além das possibilidades elencadas no caput o inadimplemento de quaisquer das cláusulas estabelecidas neste instrumento, facultará a parte que não lhe deu causa, impor sua rescisão cumulada com ressarcimento de eventuais perdas e danos e lucros cessantes e multa pecuniária irredutível e não compensatória, no valor de {multa_formatada} ({multa_extenso}).")
            
            # CLÁUSULA OITAVA - FORO
            doc.add_heading('CLÁUSULA OITAVA - FORO', level=1)
            doc.add_paragraph("Elegem as partes o foro da Comarca de Belo Horizonte, Estado de Minas Gerais, para nele serem dirimidas todas e quaisquer dúvidas ou questões oriundas do presente contrato, renunciando as partes a qualquer outro, por mais especial e privilegiado que seja.")
            
            # ENCERRAMENTO
            doc.add_paragraph("E por estarem assim justos e contratados, assinam o presente em duas (02) vias de igual teor e forma, na presença de duas testemunhas, obrigando-se por si e seus sucessores, para que produzam todos os efeitos de direito.")
            
            # Local e data
            p_data = doc.add_paragraph()
            p_data.paragraph_format.space_before = Pt(10)
            p_data.paragraph_format.space_after = Pt(10)
            p_data.add_run(f"Belo Horizonte -- MG, {data_extenso}.")
            
            # CORREÇÃO: ASSINATURAS - sem quebra de linha, com espaçamento adequado, ALINHADAS À ESQUERDA
            # Linha contratante
            p_linha1 = doc.add_paragraph()
            p_linha1.paragraph_format.space_before = Pt(30)  # Aumentado para permitir assinatura digital
            p_linha1.add_run("_" * 60)  # Reduzido para não quebrar linha
            
            p_nome_cliente = doc.add_paragraph()
            p_nome_cliente.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ALINHADO À ESQUERDA
            run_nome_cliente = p_nome_cliente.add_run(dados_cliente['nome'])
            run_nome_cliente.bold = True
            
            # Linha contratada
            p_linha2 = doc.add_paragraph()
            p_linha2.paragraph_format.space_before = Pt(30)  # Aumentado para permitir assinatura digital
            p_linha2.add_run("_" * 60)
            
            p_nome_fornecedor = doc.add_paragraph()
            p_nome_fornecedor.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ALINHADO À ESQUERDA
            run_nome_fornecedor = p_nome_fornecedor.add_run(dados_fornecedor['nome'])
            run_nome_fornecedor.bold = True
            
            # CORREÇÃO: TESTEMUNHAS - com separação entre as duas linhas
            p_test = doc.add_paragraph()
            p_test.paragraph_format.space_before = Pt(20)
            run_test = p_test.add_run("Testemunhas:")
            run_test.bold = True
            
            # Primeira testemunha
            p_linha_test1 = doc.add_paragraph()
            p_linha_test1.paragraph_format.space_before = Pt(10)
            p_linha_test1.add_run("_" * 60)
            
            p_rg1 = doc.add_paragraph()
            p_rg1.add_run("RG n.º ")
            
            # Segunda testemunha
            p_linha_test2 = doc.add_paragraph()
            p_linha_test2.paragraph_format.space_before = Pt(15)  # Separação entre testemunhas
            p_linha_test2.add_run("_" * 60)
            
            p_rg2 = doc.add_paragraph()
            p_rg2.add_run("RG n.º ")
            
            # CORREÇÃO: DADOS BANCÁRIOS - com espaçamento aumentado
            p_dados_banc = doc.add_paragraph()
            p_dados_banc.paragraph_format.space_before = Pt(50)  # Aumentado
            run_dados_banc = p_dados_banc.add_run("DADOS BANCÁRIOS PARA PAGAMENTO DA PRESTAÇÃO DE SERVIÇOS:")
            run_dados_banc.bold = True
            
            p_fornec_banco = doc.add_paragraph()
            run_fornec_banco = p_fornec_banco.add_run(dados_fornecedor['nome'])
            run_fornec_banco.bold = True
            
            doc.add_paragraph(dados_bancarios)
            
            # Salvar arquivo
            nome_arquivo = f"Contrato_{nome_cliente.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            arquivo_saida = self.PASTA_CONTRATOS / nome_arquivo
            
            doc.save(str(arquivo_saida))
            
            logger.info(f"✅ Contrato gerado com sucesso: {arquivo_saida}")
            return str(arquivo_saida)
            
        except Exception as e:
            logger.error(f"Erro ao gerar contrato: {e}")
            import traceback
            traceback.print_exc()
            return None


if __name__ == "__main__":
    # Teste básico
    print("Testando GeradorContrato com python-docx...")
    gerador = GeradorContrato()
    print(f"Categorias disponíveis: {len(gerador.listar_categorias_servicos())}")