# -*- coding: utf-8 -*-
"""
Verificador Pré-Build - Sistema de Gestão Financeira
Execute ANTES do build para garantir que tudo está OK
"""

import sys
import os
from pathlib import Path

def print_header(texto):
    print("\n" + "=" * 70)
    print(texto.center(70))
    print("=" * 70)

def print_ok(texto):
    print(f"   ✅ {texto}")

def print_erro(texto):
    print(f"   ❌ {texto}")

def print_aviso(texto):
    print(f"   ⚠️  {texto}")

def verificar_python():
    """Verifica versão do Python"""
    print("\n🐍 PYTHON")
    
    version = sys.version_info
    versao_str = f"{version.major}.{version.minor}.{version.micro}"
    
    print_ok(f"Python {versao_str}")
    
    if version.major == 3 and version.minor >= 8:
        print_ok("Versão compatível")
        return True
    else:
        print_erro("Python 3.8+ necessário")
        return False

def verificar_dependencias():
    """Verifica todas as dependências"""
    print("\n📦 DEPENDÊNCIAS")
    
    dependencias = {
        'docx': 'python-docx',
        'lxml': 'lxml',
        'lxml.etree': 'lxml.etree',
        'openpyxl': 'openpyxl',
        'pandas': 'pandas',
        'PIL': 'Pillow',
        'reportlab': 'reportlab',
        'matplotlib': 'matplotlib',
        'xlwings': 'xlwings',
        'tkcalendar': 'tkcalendar',
        'babel': 'babel',
        'num2words': 'num2words',
        'dotenv': 'python-dotenv',
    }
    
    todas_ok = True
    faltando = []
    
    for modulo, nome_pip in dependencias.items():
        try:
            mod = __import__(modulo)
            
            # Tentar pegar versão
            versao = "instalado"
            if hasattr(mod, '__version__'):
                versao = mod.__version__
            
            print_ok(f"{nome_pip:20s} → {versao}")
            
        except ImportError:
            print_erro(f"{nome_pip:20s} → NÃO INSTALADO")
            faltando.append(nome_pip)
            todas_ok = False
    
    # PyInstaller
    try:
        import PyInstaller
        print_ok(f"{'pyinstaller':20s} → {PyInstaller.__version__}")
    except ImportError:
        print_erro(f"{'pyinstaller':20s} → NÃO INSTALADO")
        faltando.append('pyinstaller')
        todas_ok = False
    
    if faltando:
        print("\n❌ INSTALE AS DEPENDÊNCIAS FALTANDO:")
        print(f"\n   pip install {' '.join(faltando)}\n")
    
    return todas_ok

def verificar_estrutura():
    """Verifica estrutura de arquivos do projeto"""
    print("\n📁 ESTRUTURA DO PROJETO")
    
    arquivos_importantes = [
        ("src/sistema_principal.py", True),
        ("src/gestao_medicoes.py", True),
        ("src/modules/gerador_contrato.py", True),
        ("src/relatorios_interface.py", True),
        ("src/ambiente_config.py", True),
        ("logo.png", False),
        ("logo1.png", False),
        ("build_sistema.spec", True),
        ("build_otimizado.py", True),
    ]
    
    todos_ok = True
    
    for arquivo, obrigatorio in arquivos_importantes:
        existe = os.path.exists(arquivo)
        
        if existe:
            print_ok(arquivo)
        else:
            if obrigatorio:
                print_erro(f"{arquivo} - OBRIGATÓRIO")
                todos_ok = False
            else:
                print_aviso(f"{arquivo} - opcional")
    
    return todos_ok

def testar_import_docx():
    """Testa import específico do python-docx"""
    print("\n🔬 TESTE ESPECÍFICO: python-docx")
    
    try:
        import docx
        print_ok("import docx")
        
        from docx import Document
        print_ok("from docx import Document")
        
        from docx.shared import Pt, Inches
        print_ok("from docx.shared import Pt, Inches")
        
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print_ok("from docx.enum.text import WD_ALIGN_PARAGRAPH")
        
        # Testar criação de documento
        doc = Document()
        print_ok("Document() criado com sucesso")
        
        # Testar adição de parágrafo
        doc.add_paragraph("Teste")
        print_ok("add_paragraph() funciona")
        
        print("\n✅ python-docx está 100% funcional!")
        return True
        
    except Exception as e:
        print_erro(f"Erro ao testar python-docx: {e}")
        return False

def testar_import_lxml():
    """Testa import do lxml"""
    print("\n🔬 TESTE ESPECÍFICO: lxml")
    
    try:
        import lxml
        print_ok("import lxml")
        
        from lxml import etree
        print_ok("from lxml import etree")
        
        # Testar criação de elemento
        root = etree.Element("root")
        print_ok("etree.Element() funciona")
        
        print("\n✅ lxml está 100% funcional!")
        return True
        
    except Exception as e:
        print_erro(f"Erro ao testar lxml: {e}")
        return False

def verificar_arquivo_spec():
    """Verifica se o arquivo .spec está correto"""
    print("\n⚙️  ARQUIVO SPEC")
    
    if not os.path.exists('build_sistema.spec'):
        print_erro("build_sistema.spec não encontrado!")
        print("   Execute novamente o script que gera este arquivo.")
        return False
    
    try:
        with open('build_sistema.spec', 'r', encoding='utf-8') as f:
            conteudo = f.read()
        
        # Verificar elementos importantes
        checks = [
            ("collect_all('docx')", "Coleta automática do python-docx"),
            ("collect_all('lxml')", "Coleta automática do lxml"),
            ("NOME_EXECUTAVEL", "Nome do executável definido"),
            ("sistema_principal.py", "Arquivo principal definido"),
        ]
        
        todos_ok = True
        for texto, descricao in checks:
            if texto in conteudo:
                print_ok(descricao)
            else:
                print_erro(f"{descricao} - NÃO ENCONTRADO")
                todos_ok = False
        
        return todos_ok
        
    except Exception as e:
        print_erro(f"Erro ao ler .spec: {e}")
        return False

def resumo_final(resultados):
    """Mostra resumo final"""
    print_header("RESUMO")
    
    total = len(resultados)
    ok = sum(resultados.values())
    
    print(f"\n   Verificações realizadas: {total}")
    print(f"   ✅ Aprovadas: {ok}")
    print(f"   ❌ Reprovadas: {total - ok}")
    
    if all(resultados.values()):
        print("\n" + "=" * 70)
        print("✅ TUDO OK! PODE EXECUTAR O BUILD".center(70))
        print("=" * 70)
        print("\nPróximo passo:")
        print("   python build_otimizado.py")
        print("=" * 70)
        return True
    else:
        print("\n" + "=" * 70)
        print("❌ CORRIJA OS PROBLEMAS ANTES DO BUILD".center(70))
        print("=" * 70)
        
        print("\nProblemas encontrados:")
        for nome, passou in resultados.items():
            if not passou:
                print(f"   ❌ {nome}")
        
        print("\n" + "=" * 70)
        return False

def main():
    """Função principal"""
    print_header("VERIFICADOR PRÉ-BUILD")
    print("Sistema de Gestão Financeira")
    print("\nEste script verifica se tudo está OK ANTES do build.")
    
    resultados = {}
    
    # Verificações
    resultados['Python'] = verificar_python()
    resultados['Dependências'] = verificar_dependencias()
    resultados['Estrutura'] = verificar_estrutura()
    resultados['python-docx'] = testar_import_docx()
    resultados['lxml'] = testar_import_lxml()
    resultados['Arquivo SPEC'] = verificar_arquivo_spec()
    
    # Resumo
    sucesso = resumo_final(resultados)
    
    return 0 if sucesso else 1

if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n\nVerificação cancelada")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Erro inesperado: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
